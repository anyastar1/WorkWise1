"""
Утилиты для обработки документов (PDF/DOCX -> PNG)
"""

import os
import hashlib
import zipfile
import io
from pathlib import Path
from werkzeug.utils import secure_filename
from database import Document, Page


def calculate_file_hash(file_content: bytes) -> str:
    """Вычисляет SHA-256 хеш файла"""
    return hashlib.sha256(file_content).hexdigest()


def validate_file_type(file_content: bytes, filename: str) -> tuple[bool, str]:
    """
    Проверяет тип файла по сигнатурам и расширению
    
    Returns:
        (is_valid, file_type) где file_type: 'PDF' или 'DOCX'
    """
    # Проверка по расширению
    ext = filename.rsplit(".", 1)[1].lower() if "." in filename else ""
    
    # Проверка PDF по сигнатуре (PDF начинается с %PDF)
    if ext == "pdf":
        if file_content[:4] == b"%PDF":
            return True, "PDF"
    
    # Проверка DOCX по сигнатуре (DOCX - это ZIP архив с определенной структурой)
    if ext == "docx":
        # DOCX файлы начинаются с PK (ZIP signature)
        if file_content[:2] == b"PK":
            # Проверяем наличие типичных файлов DOCX в ZIP
            try:
                zip_file = zipfile.ZipFile(io.BytesIO(file_content))
                # DOCX должен содержать файл [Content_Types].xml
                if "[Content_Types].xml" in zip_file.namelist():
                    return True, "DOCX"
            except:
                pass
    
    return False, ""


def process_document(file, user_id: int, operation_type: str, db_session) -> Document:
    """
    Обрабатывает загруженный документ:
    1. Сохраняет файл
    2. Конвертирует в изображения
    3. Сохраняет в БД
    
    Args:
        file: Файловый объект из request.files
        user_id: ID пользователя
        operation_type: Тип операции
        db_session: Сессия БД
        
    Returns:
        Document объект
    """
    # Читаем содержимое файла
    file_content = file.read()
    file.seek(0)
    
    # Вычисляем хеш
    file_hash = calculate_file_hash(file_content)
    
    # Проверяем MIME-тип
    is_valid, file_type = validate_file_type(file_content, file.filename)
    if not is_valid:
        raise ValueError("Недопустимый формат файла. Разрешены только PDF и DOCX.")
    
    # Проверяем, существует ли уже документ с таким хешем
    existing_doc = db_session.query(Document).filter_by(hash=file_hash).first()
    if existing_doc:
        # Если документ уже существует, возвращаем его
        return existing_doc
    
    # Определяем расширение
    ext = "pdf" if file_type == "PDF" else "docx"
    
    # Создаем структуру директорий
    doc_dir = Path("uploads") / file_hash
    pages_dir = doc_dir / "pages"
    pages_dir.mkdir(parents=True, exist_ok=True)
    
    # Сохраняем оригинальный файл
    original_path = doc_dir / f"original.{ext}"
    with open(original_path, "wb") as f:
        f.write(file_content)
    
    # Конвертируем в изображения
    page_paths = convert_to_images(original_path, pages_dir, file_type)
    
    # Создаем запись в БД
    document = Document(
        user_id=user_id,
        file_type=file_type,
        operation_type=operation_type,
        hash=file_hash
    )
    db_session.add(document)
    db_session.flush()  # Получаем ID документа
    
    # Создаем записи страниц
    for page_num, page_path in enumerate(page_paths, start=1):
        page = Page(
            document_id=document.id,
            page_number=page_num,
            image_path=str(page_path.relative_to("uploads"))
        )
        db_session.add(page)
    
    db_session.commit()
    
    return document


def convert_to_images(file_path: Path, output_dir: Path, file_type: str) -> list[Path]:
    """
    Конвертирует документ в изображения PNG
    
    Args:
        file_path: Путь к файлу
        output_dir: Директория для сохранения изображений
        file_type: Тип файла ('PDF' или 'DOCX')
        
    Returns:
        Список путей к созданным изображениям
    """
    page_paths = []
    
    if file_type == "PDF":
        page_paths = convert_pdf_to_images(file_path, output_dir)
    elif file_type == "DOCX":
        page_paths = convert_docx_to_images(file_path, output_dir)
    
    return page_paths


def convert_pdf_to_images(pdf_path: Path, output_dir: Path, dpi: int = 70) -> list[Path]:
    """Конвертирует PDF в PNG изображения"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF не установлен. Установите: pip install PyMuPDF")
    
    page_paths = []
    pdf_document = fitz.open(pdf_path)
    
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        
        # Конвертируем страницу в изображение с разрешением 70 DPI
        mat = fitz.Matrix(dpi / 72, dpi / 72)  # Масштабирование для DPI
        pix = page.get_pixmap(matrix=mat)
        
        # Сохраняем как PNG
        page_path = output_dir / f"page_{page_num + 1}.png"
        pix.save(page_path)
        page_paths.append(page_path)
    
    pdf_document.close()
    return page_paths


def convert_docx_to_images(docx_path: Path, output_dir: Path, dpi: int = 70) -> list[Path]:
    """Конвертирует DOCX в PNG изображения"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF не установлен. Установите: pip install PyMuPDF")
    
    # Пробуем использовать docx2pdf, если доступен
    temp_pdf_path = output_dir / "temp.pdf"
    page_paths = []
    
    try:
        # Попытка использовать docx2pdf (требует LibreOffice)
        try:
            from docx2pdf import convert as docx2pdf_convert
            docx2pdf_convert(str(docx_path), str(temp_pdf_path))
            
            # Конвертируем PDF в изображения
            page_paths = convert_pdf_to_images(temp_pdf_path, output_dir, dpi)
            
            # Удаляем временный PDF
            if temp_pdf_path.exists():
                temp_pdf_path.unlink()
            
            return page_paths
        except ImportError:
            # docx2pdf не установлен, используем альтернативный метод
            pass
        except Exception as e:
            # docx2pdf не сработал (возможно, нет LibreOffice)
            # Пробуем альтернативный метод
            pass
        
        # Альтернативный метод: конвертируем DOCX напрямую через python-docx и PIL
        # Это менее точный метод, но не требует внешних зависимостей
        from docx import Document as DocxDocument
        from PIL import Image, ImageDraw, ImageFont
        import io
        
        doc = DocxDocument(str(docx_path))
        
        # Для каждой страницы создаем изображение
        # Примечание: python-docx не поддерживает разбиение на страницы напрямую,
        # поэтому создаем одно изображение для всего документа
        # В реальном приложении лучше использовать docx2pdf или другой инструмент
        
        # Простой подход: создаем изображение текста
        # Это упрощенная версия, в продакшене лучше использовать docx2pdf
        
        # Если docx2pdf недоступен, создаем одно изображение со всем содержимым
        if not page_paths:
            # Создаем изображение с текстом документа
            img_width = 800
            img_height = 1000
            
            img = Image.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 16)
            except:
                font = ImageFont.load_default()
            
            y_position = 50
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Простое отображение текста
                    draw.text((50, y_position), paragraph.text[:100], fill='black', font=font)
                    y_position += 30
                    if y_position > img_height - 50:
                        break
            
            page_path = output_dir / "page_1.png"
            img.save(page_path)
            page_paths.append(page_path)
        
        return page_paths if page_paths else [output_dir / "page_1.png"]
        
    except Exception as e:
        raise Exception(f"Ошибка конвертации DOCX: {str(e)}. Убедитесь, что установлены необходимые библиотеки.")
