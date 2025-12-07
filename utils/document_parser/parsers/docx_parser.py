"""
DOCX Parser using python-docx.

Extracts text with approximate positioning and styles from DOCX documents.

NOTE: DOCX is a flow format without fixed coordinates.
Coordinates are calculated approximately based on styles.
"""

from docx import Document as DocxDocument
from docx.shared import Pt
from pathlib import Path
from typing import Union, BinaryIO, List, Optional
import hashlib

from .base import BaseParser
from src.models.document import (
    ParsedDocument,
    DocumentMetadata,
    DocumentType,
    DocumentPage,
    PageInfo,
    TextBlock,
    TextBlockType,
    TextLine,
    TextSpan,
    TextStyle,
    BoundingBox,
)


class DOCXParser(BaseParser):
    """
    DOCX document parser.
    
    IMPORTANT: DOCX is a flow format without fixed coordinates.
    Coordinates are calculated approximately based on styles.
    
    For accurate coordinates, convert DOCX to PDF first.
    """
    
    def __init__(self, 
                 page_width: float = 595.0, 
                 page_height: float = 842.0,
                 margin_left: float = 72.0, 
                 margin_top: float = 72.0):
        """
        Initialize DOCX parser.
        
        Args:
            page_width: Page width in pt (A4 = 595)
            page_height: Page height in pt (A4 = 842)
            margin_left: Left margin in pt
            margin_top: Top margin in pt
        """
        self.page_width = page_width
        self.page_height = page_height
        self.margin_left = margin_left
        self.margin_top = margin_top
        self.content_width = page_width - 2 * margin_left
        self._block_counter = 0
    
    def parse(self, source: Union[str, Path, BinaryIO]) -> ParsedDocument:
        """
        Parse a DOCX document.
        
        Args:
            source: Path to DOCX file or binary stream
            
        Returns:
            ParsedDocument with extracted content
        """
        self._block_counter = 0
        doc = DocxDocument(source)
        
        metadata = self._extract_metadata(doc, source)
        pages = self._extract_pages(doc)
        
        # Update total_pages in metadata
        metadata.total_pages = len(pages)
        
        return ParsedDocument(metadata=metadata, pages=pages)
    
    def _extract_metadata(self, doc: DocxDocument, source) -> DocumentMetadata:
        """Extract document metadata."""
        filepath = Path(source) if isinstance(source, (str, Path)) else None
        core_props = doc.core_properties
        
        # Calculate content hash
        content_hash = None
        if filepath and filepath.exists():
            with open(filepath, 'rb') as f:
                content_hash = hashlib.sha256(f.read()).hexdigest()[:16]
        
        return DocumentMetadata(
            filename=filepath.name if filepath else "unknown.docx",
            document_type=DocumentType.DOCX,
            total_pages=1,  # Will be updated after parsing
            file_size_bytes=filepath.stat().st_size if filepath and filepath.exists() else 0,
            title=core_props.title,
            author=core_props.author,
            creation_date=str(core_props.created) if core_props.created else None,
            modification_date=str(core_props.modified) if core_props.modified else None,
            content_hash=content_hash
        )
    
    def _extract_pages(self, doc: DocxDocument) -> List[DocumentPage]:
        """
        Extract content with approximate coordinate calculations.
        
        Page breaks are determined by height.
        """
        pages = []
        current_y = self.margin_top
        current_page_num = 1
        current_blocks = []
        reading_order = 0
        
        for para in doc.paragraphs:
            block = self._parse_paragraph(para, current_page_num, reading_order, current_y)
            
            if block is None:
                continue
            
            # Check if block fits on current page
            block_height = block.bbox.height
            if current_y + block_height > self.page_height - self.margin_top:
                # Save current page
                if current_blocks:
                    pages.append(DocumentPage(
                        info=PageInfo(
                            page_number=current_page_num,
                            width=self.page_width,
                            height=self.page_height
                        ),
                        blocks=current_blocks
                    ))
                
                # Start new page
                current_page_num += 1
                current_y = self.margin_top
                current_blocks = []
                
                # Update block coordinates for new page
                block = self._parse_paragraph(para, current_page_num, reading_order, current_y)
            
            if block:
                current_blocks.append(block)
                current_y = block.bbox.y1 + 10  # Paragraph spacing
                reading_order += 1
        
        # Add last page
        if current_blocks:
            pages.append(DocumentPage(
                info=PageInfo(
                    page_number=current_page_num,
                    width=self.page_width,
                    height=self.page_height
                ),
                blocks=current_blocks
            ))
        
        return pages if pages else [DocumentPage(
            info=PageInfo(page_number=1, width=self.page_width, height=self.page_height),
            blocks=[]
        )]
    
    def _parse_paragraph(self, para, page_num: int, order: int, y_offset: float) -> Optional[TextBlock]:
        """Parse a DOCX paragraph."""
        text = para.text.strip()
        if not text:
            return None
        
        self._block_counter += 1
        block_id = f"blk_p{page_num}_{self._block_counter}"
        
        # Determine paragraph style
        style = para.style
        font_size = self._get_font_size(para) or 12.0
        
        # Calculate approximate block height
        line_height = font_size * 1.2
        estimated_lines = max(1, len(text) / 80)  # ~80 chars per line
        block_height = line_height * estimated_lines
        
        # Determine block type
        block_type = self._detect_block_type(para, style)
        
        bbox = BoundingBox(
            x0=self.margin_left,
            y0=y_offset,
            x1=self.margin_left + self.content_width,
            y1=y_offset + block_height
        )
        
        # Create lines (simplified - entire text as one line)
        lines = self._create_lines(para, bbox)
        
        return TextBlock(
            block_id=block_id,
            block_type=block_type,
            lines=lines,
            bbox=bbox,
            page_number=page_num,
            reading_order=order,
            semantic_level=self._get_heading_level(style) if block_type == TextBlockType.HEADING else None
        )
    
    def _create_lines(self, para, block_bbox: BoundingBox) -> List[TextLine]:
        """Create lines from paragraph runs."""
        spans = []
        current_x = block_bbox.x0
        
        for run in para.runs:
            if not run.text:
                continue
            
            span = self._parse_run(run, current_x, block_bbox.y0)
            if span:
                spans.append(span)
                current_x = span.bbox.x1
        
        if not spans:
            return []
        
        # All spans in one line (simplification)
        line_bbox = BoundingBox(
            x0=block_bbox.x0,
            y0=block_bbox.y0,
            x1=max(s.bbox.x1 for s in spans),
            y1=block_bbox.y1
        )
        
        return [TextLine(spans=spans, bbox=line_bbox)]
    
    def _parse_run(self, run, x_offset: float, y_offset: float) -> Optional[TextSpan]:
        """Parse a run (fragment with uniform formatting)."""
        text = run.text
        if not text:
            return None
        
        font = run.font
        font_size = float(font.size.pt) if font.size else 12.0
        
        # Approximate text width
        char_width = font_size * 0.5
        text_width = len(text) * char_width
        
        # Color
        color = "#000000"
        if font.color and font.color.rgb:
            color = f"#{font.color.rgb}"
        
        style = TextStyle(
            font_name=font.name or "default",
            font_size=font_size,
            font_weight="bold" if font.bold else "normal",
            font_style="italic" if font.italic else "normal",
            color=color,
            is_underline=font.underline or False,
            is_strikethrough=font.strike or False
        )
        
        bbox = BoundingBox(
            x0=x_offset,
            y0=y_offset,
            x1=x_offset + text_width,
            y1=y_offset + font_size * 1.2
        )
        
        return TextSpan(text=text, style=style, bbox=bbox)
    
    def _get_font_size(self, para) -> Optional[float]:
        """Get paragraph font size."""
        for run in para.runs:
            if run.font.size:
                return float(run.font.size.pt)
        
        # From style
        if para.style and para.style.font and para.style.font.size:
            return float(para.style.font.size.pt)
        
        return None
    
    def _detect_block_type(self, para, style) -> TextBlockType:
        """Determine block type."""
        style_name = style.name.lower() if style else ""
        
        if "heading" in style_name or "title" in style_name:
            return TextBlockType.HEADING
        
        if "list" in style_name or para.text.strip().startswith(("•", "-", "–", "●", "○")):
            return TextBlockType.LIST_ITEM
        
        return TextBlockType.PARAGRAPH
    
    def _get_heading_level(self, style) -> Optional[int]:
        """Extract heading level from style."""
        if not style:
            return None
        
        name = style.name.lower()
        for i in range(1, 7):
            if f"heading {i}" in name or f"heading{i}" in name:
                return i
        
        if "title" in name:
            return 1
        
        return None
    
    def supports_format(self, filepath: Union[str, Path]) -> bool:
        """Check if DOCX format is supported."""
        return str(filepath).lower().endswith('.docx')
